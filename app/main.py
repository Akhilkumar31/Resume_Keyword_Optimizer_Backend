"""FastAPI application entry point for Resume Keyword Optimizer."""

from fastapi import FastAPI, HTTPException, File, UploadFile, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, PlainTextResponse
from typing import Optional
import logging
import requests

from app.schemas import (
    ResumeAnalysisRequest, 
    ResumeAnalysis,
    KeywordExtractResponse,
    JobDescriptionAnalysisResponse,
    ComparisonMetrics,
    FetchJobDescriptionRequest,
    FetchJobDescriptionResponse,
    ReportRequest
)
from app.parser import ResumeParser
from app.analyzer import KeywordAnalyzer
from app.scraper import scrape_job_description

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Constants
MAX_UPLOAD_SIZE = 5 * 1024 * 1024  # 5MB
SUPPORTED_FILE_FORMATS = ('.txt', '.pdf', '.docx')
DEFAULT_TOP_N_KEYWORDS = 20
DEFAULT_JD_TOP_N = 30

# Initialize FastAPI app
app = FastAPI(
    title="Resume Keyword Optimizer API",
    description="API for optimizing resumes with keyword analysis",
    version="1.0.0"
)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, specify allowed origins
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize components
resume_parser = ResumeParser()
keyword_analyzer = KeywordAnalyzer()


@app.get("/", tags=["Health"])
async def root():
    """Root endpoint - API is running."""
    return {
        "message": "Resume Keyword Optimizer API",
        "version": "1.0.0",
        "status": "running"
    }


@app.get("/health", tags=["Health"])
async def health_check():
    """Health check endpoint."""
    return {"status": "healthy"}


@app.post("/analyze", response_model=ResumeAnalysis, tags=["Analysis"])
async def analyze_resume(request: ResumeAnalysisRequest) -> ResumeAnalysis:
    """
    Analyze a resume and optionally match against a job description.
    
    Args:
        request: ResumeAnalysisRequest containing resume text and optional job description
        
    Returns:
        ResumeAnalysis: Analysis results with keyword matching and recommendations
        
    Raises:
        HTTPException: If resume text is empty or job description is empty when provided
    """
    try:
        # Validate resume text
        if not request.resume_text or not request.resume_text.strip():
            raise HTTPException(status_code=400, detail="Resume text cannot be empty")
        
        # Validate job description if provided
        if request.job_description is not None:
            if not request.job_description.strip():
                raise HTTPException(status_code=400, detail="Job description cannot be empty if provided")
        
        logger.info("Analyzing resume")
        
        # Analyze the resume
        analysis = keyword_analyzer.analyze_resume(
            request.resume_text,
            request.job_description if request.job_description and request.job_description.strip() else None
        )
        
        return analysis
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error analyzing resume: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/analyze/file", response_model=ResumeAnalysis, tags=["Analysis"])
async def analyze_resume_file(
    resume: UploadFile = File(..., description="Resume file (.txt, .pdf, or .docx)"),
    job_description: Optional[str] = Form(None, description="Optional job description text")
) -> ResumeAnalysis:
    """
    Analyze a resume file and optionally match against a job description.
    
    Accepts resume as a file upload (.txt, .pdf, .docx) and job description as form data.
    
    Args:
        resume: Resume file to analyze (required)
        job_description: Optional job description text for matching
        
    Returns:
        ResumeAnalysis: Analysis results with score, matched keywords, missing keywords, and suggestions
        
    Raises:
        HTTPException: If resume file is invalid, empty, or unsupported format
        HTTPException: If job description is provided but empty
    """
    try:
        # Validate resume file exists
        if not resume:
            raise HTTPException(status_code=400, detail="Resume file is required")
        
        if not resume.filename:
            raise HTTPException(status_code=400, detail="No file provided")
        
        # Check file extension
        file_extension = resume.filename.lower().split('.')[-1]
        if f".{file_extension}" not in SUPPORTED_FILE_FORMATS:
            raise HTTPException(
                status_code=400,
                detail=f"Only {', '.join(SUPPORTED_FILE_FORMATS)} files are supported. Got: .{file_extension}"
            )
        
        # Read file content
        content = await resume.read()
        
        # Check file size
        if len(content) > MAX_UPLOAD_SIZE:
            raise HTTPException(
                status_code=413,
                detail="File size exceeds maximum allowed size of 5MB"
            )
        
        # Extract resume text based on file format
        if file_extension == 'txt':
            try:
                resume_text = content.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    resume_text = content.decode('latin-1')
                except UnicodeDecodeError:
                    raise HTTPException(
                        status_code=400,
                        detail="Text file must be valid UTF-8 or Latin-1 encoded"
                    )
        elif file_extension == 'pdf':
            try:
                import pdfplumber
                import tempfile
                # Write to temporary file for pdfplumber
                with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                    tmp.write(content)
                    tmp_path = tmp.name
                
                text = ""
                with pdfplumber.open(tmp_path) as pdf:
                    if not pdf.pages:
                        raise HTTPException(status_code=400, detail="PDF file is empty or cannot be read")
                    for page in pdf.pages:
                        extracted_text = page.extract_text()
                        if extracted_text:
                            text += extracted_text + "\n"
                
                resume_text = text.strip()
                
                # Clean up temp file
                import os as os_module
                os_module.unlink(tmp_path)
                
                if not resume_text:
                    raise HTTPException(status_code=400, detail="PDF file contains no readable text")
            except HTTPException:
                raise
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Error reading PDF file: {str(e)}")
        elif file_extension == 'docx':
            try:
                from docx import Document
                import tempfile
                # Write to temporary file for python-docx
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                    tmp.write(content)
                    tmp_path = tmp.name
                
                doc = Document(tmp_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                
                resume_text = text.strip()
                
                # Clean up temp file
                import os as os_module
                os_module.unlink(tmp_path)
                
                if not resume_text:
                    raise HTTPException(status_code=400, detail="DOCX file contains no readable text")
            except HTTPException:
                raise
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Error reading DOCX file: {str(e)}")
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file format: .{file_extension}")
        
        # Validate resume text
        if not resume_text.strip():
            raise HTTPException(status_code=400, detail="Resume file is empty or contains no readable text")
        
        # Validate job description if provided
        if job_description is not None and not job_description.strip():
            raise HTTPException(status_code=400, detail="Job description cannot be empty if provided")
        
        logger.info(f"Analyzing resume file: {resume.filename}")
        
        # Analyze the resume
        analysis = keyword_analyzer.analyze_resume(
            resume_text,
            job_description if job_description and job_description.strip() else None
        )
        
        return analysis
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error analyzing resume file: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")


@app.post("/parse", tags=["Parsing"])
async def parse_resume(request: ResumeAnalysisRequest):
    """
    Parse a resume and extract structured information.
    
    Args:
        request: ResumeAnalysisRequest containing resume text
        
    Returns:
        ParsedResume: Structured resume data
        
    Raises:
        HTTPException: If resume text is empty
    """
    try:
        if not request.resume_text.strip():
            raise HTTPException(status_code=400, detail="Resume text cannot be empty")
        
        logger.info("Parsing resume")
        
        # Parse the resume
        parsed = resume_parser.parse(request.resume_text)
        
        return parsed
    
    except Exception as e:
        logger.error(f"Error parsing resume: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/upload", tags=["File Upload"])
async def upload_resume(file: UploadFile = File(...)):
    """
    Upload and process a resume file.
    
    Supported formats: .txt, .pdf, .docx
    Maximum file size: 5MB
    
    Args:
        file: Resume file to upload
        
    Returns:
        Parsed resume data with sections, contact info, skills, experience, and education
        
    Raises:
        HTTPException: If file type is not supported or file is too large or empty
    """
    try:
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")
        
        # Get file extension
        file_extension = file.filename.lower().split('.')[-1]
        
        # Check file type
        if f".{file_extension}" not in SUPPORTED_FILE_FORMATS:
            raise HTTPException(
                status_code=400,
                detail=f"Only {', '.join(SUPPORTED_FILE_FORMATS)} files are supported. Got: .{file_extension}"
            )
        
        # Read file content
        content = await file.read()
        
        # Check file size
        if len(content) > MAX_UPLOAD_SIZE:
            raise HTTPException(
                status_code=413,
                detail="File size exceeds maximum allowed size of 5MB"
            )
        
        # Extract resume text based on file format
        if file_extension == 'txt':
            try:
                text_content = content.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    text_content = content.decode('latin-1')
                except UnicodeDecodeError:
                    raise HTTPException(
                        status_code=400,
                        detail="Text file must be valid UTF-8 or Latin-1 encoded"
                    )
        elif file_extension == 'pdf':
            try:
                import pdfplumber
                import tempfile
                # Write to temporary file for pdfplumber
                with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
                    tmp.write(content)
                    tmp_path = tmp.name
                
                text = ""
                with pdfplumber.open(tmp_path) as pdf:
                    if not pdf.pages:
                        raise HTTPException(status_code=400, detail="PDF file is empty or cannot be read")
                    for page in pdf.pages:
                        extracted_text = page.extract_text()
                        if extracted_text:
                            text += extracted_text + "\n"
                
                text_content = text.strip()
                
                # Clean up temp file
                import os as os_module
                os_module.unlink(tmp_path)
                
                if not text_content:
                    raise HTTPException(status_code=400, detail="PDF file contains no readable text")
            except HTTPException:
                raise
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Error reading PDF file: {str(e)}")
        elif file_extension == 'docx':
            try:
                from docx import Document
                import tempfile
                # Write to temporary file for python-docx
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                    tmp.write(content)
                    tmp_path = tmp.name
                
                doc = Document(tmp_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                
                text_content = text.strip()
                
                # Clean up temp file
                import os as os_module
                os_module.unlink(tmp_path)
                
                if not text_content:
                    raise HTTPException(status_code=400, detail="DOCX file contains no readable text")
            except HTTPException:
                raise
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Error reading DOCX file: {str(e)}")
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file format: .{file_extension}")
        
        # Validate text content
        if not text_content or not text_content.strip():
            raise HTTPException(status_code=400, detail="File is empty or contains no readable text")
        
        logger.info(f"Uploading resume file: {file.filename}")
        
        # Parse the resume
        parsed = resume_parser.parse(text_content)
        
        return parsed
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error uploading resume: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/keywords/extract", response_model=KeywordExtractResponse, tags=["Keywords"])
async def extract_keywords(request: ResumeAnalysisRequest) -> KeywordExtractResponse:
    """
    Extract top keywords from text (resume or job description).
    
    Args:
        request: ResumeAnalysisRequest with resume_text containing the text to analyze
        
    Returns:
        KeywordExtractResponse with top keywords and their frequencies
        
    Raises:
        HTTPException: If text is empty
    """
    try:
        if not request.resume_text.strip():
            raise HTTPException(status_code=400, detail="Text cannot be empty")
        
        logger.info("Extracting keywords from text")
        
        top_keywords = keyword_analyzer.extract_top_keywords(request.resume_text, top_n=DEFAULT_TOP_N_KEYWORDS)
        
        return KeywordExtractResponse(
            top_keywords=[{"keyword": kw, "frequency": freq} for kw, freq in top_keywords],
            total_unique_keywords=len(top_keywords)
        )
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error extracting keywords: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/keywords/analyze-jd", response_model=JobDescriptionAnalysisResponse, tags=["Keywords"])
async def analyze_job_description(request: ResumeAnalysisRequest) -> JobDescriptionAnalysisResponse:
    """
    Analyze a job description to extract key requirements.
    
    Args:
        request: ResumeAnalysisRequest with resume_text containing the job description
        
    Returns:
        JobDescriptionAnalysisResponse with top keywords and analysis
        
    Raises:
        HTTPException: If job description is empty
    """
    try:
        if not request.resume_text.strip():
            raise HTTPException(status_code=400, detail="Job description cannot be empty")
        
        logger.info("Analyzing job description")
        
        analysis = keyword_analyzer.analyze_job_description(request.resume_text, top_n=DEFAULT_JD_TOP_N)
        
        return JobDescriptionAnalysisResponse(
            top_keywords=[{"keyword": kw, "frequency": freq} for kw, freq in analysis["top_keywords"]],
            total_unique_keywords=analysis["total_keywords"],
            keyword_frequency=analysis["keyword_frequency"]
        )
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error analyzing job description: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/compare", response_model=ComparisonMetrics, tags=["Comparison"])
async def compare_resume_to_jd(request: ResumeAnalysisRequest) -> ComparisonMetrics:
    """
    Perform detailed comparison between resume and job description.
    
    Analyzes keyword matching, coverage, and provides similarity metrics.
    
    Args:
        request: ResumeAnalysisRequest with both resume_text and job_description
        
    Returns:
        ComparisonMetrics with detailed comparison including:
        - Match score (0-1 scale)
        - Matched keywords
        - Missing keywords
        - Keyword coverage percentage
        - Jaccard similarity score
        
    Raises:
        HTTPException: If either resume or job description is empty
    """
    try:
        if not request.resume_text.strip():
            raise HTTPException(status_code=400, detail="Resume text cannot be empty")
        if not request.job_description or not request.job_description.strip():
            raise HTTPException(status_code=400, detail="Job description cannot be empty")
        
        logger.info("Comparing resume to job description")
        
        comparison = keyword_analyzer.compare_resume_to_jd(
            request.resume_text,
            request.job_description
        )
        
        return ComparisonMetrics(
            match_score=comparison["match_score"],
            matched_count=comparison["matched_count"],
            total_jd_keywords=comparison["total_jd_keywords"],
            missing_count=comparison["missing_count"],
            keyword_coverage_percentage=comparison["keyword_coverage"],
            jaccard_similarity=comparison["jaccard_similarity"],
            top_resume_keywords=[{"keyword": kw, "frequency": freq} for kw, freq in comparison["top_resume_keywords"]],
            top_jd_keywords=[{"keyword": kw, "frequency": freq} for kw, freq in comparison["top_jd_keywords"]],
            matched_keywords=comparison["intersection_keywords"],
            missing_keywords=comparison["missing_keywords"]
        )
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error comparing resume to job description: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/fetch-job-description", response_model=FetchJobDescriptionResponse, tags=["Job Description"])
async def fetch_job_description(request: FetchJobDescriptionRequest) -> FetchJobDescriptionResponse:
    """
    Fetch and extract job description text from a given URL.
    
    Uses BeautifulSoup to parse the webpage and extract visible text.
    Filters out script, style, and navigation elements for clean content.
    
    Args:
        request: FetchJobDescriptionRequest containing the URL to fetch
        
    Returns:
        FetchJobDescriptionResponse with extracted job description text and page title
        
    Raises:
        HTTPException: If URL is invalid or request fails
    """
    try:
        logger.info(f"Fetching job description from URL: {request.url}")

        job_description, title = scrape_job_description(request.url)

        return FetchJobDescriptionResponse(
            url=request.url,
            job_description=job_description,
            title=title,
            status="success"
        )

    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except requests.exceptions.Timeout:
        raise HTTPException(status_code=408, detail="Request timeout: The URL took too long to respond")
    except requests.exceptions.ConnectionError:
        raise HTTPException(status_code=503, detail="Connection error: Unable to reach the URL")
    except requests.exceptions.HTTPError as e:
        status_code = e.response.status_code if e.response is not None else 502
        raise HTTPException(status_code=status_code, detail=f"HTTP error from target URL: {status_code}")
    except requests.exceptions.RequestException as e:
        raise HTTPException(status_code=400, detail=f"Request error: {str(e)}")
    except RuntimeError as e:
        raise HTTPException(status_code=422, detail=str(e))
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching job description: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")


@app.post("/download-report", tags=["Report"])
async def download_report(request: ReportRequest) -> PlainTextResponse:
    """
    Generate a downloadable plain-text resume analysis report.

    Args:
        request: ReportRequest containing analysis results

    Returns:
        PlainTextResponse with Content-Disposition header for file download
    """
    lines = [
        "=" * 60,
        "       RESUME ANALYSIS REPORT",
        "=" * 60,
        "",
        f"MATCH SCORE: {request.match_score:.1f}%",
        "",
        "-" * 60,
        "MATCHED KEYWORDS",
        "-" * 60,
    ]

    if request.matched_keywords:
        for item in request.matched_keywords:
            keyword = item.get("keyword", "")
            frequency = item.get("frequency", "")
            relevance = item.get("relevance_score", "")
            lines.append(f"  - {keyword}  (frequency: {frequency}, relevance: {relevance})")
    else:
        lines.append("  None")

    lines += [
        "",
        "-" * 60,
        "MISSING KEYWORDS",
        "-" * 60,
    ]

    if request.missing_keywords:
        for kw in request.missing_keywords:
            lines.append(f"  - {kw}")
    else:
        lines.append("  None")

    lines += [
        "",
        "-" * 60,
        "KEYWORD SUGGESTIONS",
        "-" * 60,
    ]

    if request.suggestions:
        for keyword, synonyms in request.suggestions.items():
            synonym_list = ", ".join(synonyms) if isinstance(synonyms, list) else str(synonyms)
            lines.append(f"  {keyword}: {synonym_list}")
    else:
        lines.append("  None")

    lines += [
        "",
        "-" * 60,
        "RECOMMENDATIONS",
        "-" * 60,
    ]

    if request.recommendations:
        for i, rec in enumerate(request.recommendations, start=1):
            lines.append(f"  {i}. {rec}")
    else:
        lines.append("  None")

    lines += [
        "",
        "=" * 60,
        "End of Report",
        "=" * 60,
    ]

    report_text = "\n".join(lines)

    return PlainTextResponse(
        content=report_text,
        headers={"Content-Disposition": "attachment; filename=\"resume_analysis_report.txt\""},
    )


@app.exception_handler(HTTPException)
async def http_exception_handler(request, exc):
    """Handle HTTP exceptions."""
    return JSONResponse(
        status_code=exc.status_code,
        content={"detail": exc.detail},
    )


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(
        "app.main:app",
        host="0.0.0.0",
        port=8000,
        reload=True
    )
