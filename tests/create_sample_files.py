"""Script to generate sample resume files in different formats for testing."""

from pathlib import Path
from docx import Document

# Get the tests directory
tests_dir = Path(__file__).parent

# Sample resume content
RESUME_CONTENT = """JANE SMITH
jane.smith@example.com | (555) 987-6543 | linkedin.com/in/janesmith

PROFESSIONAL SUMMARY
Experienced full-stack software engineer with 6 years of expertise in cloud architecture and distributed systems. 
Proven ability to lead technical initiatives and mentor development teams.

WORK EXPERIENCE
Principal Software Engineer - CloudTech Solutions
March 2022 - Present
• Led architecture redesign of core microservices platform serving 2M+ users
• Implemented automated deployment pipeline reducing release time by 75%
• Mentored team of 8 junior and senior engineers
• Drove adoption of Kubernetes reducing infrastructure costs by 40%

Senior Software Engineer - Digital Innovations Inc.
July 2019 - February 2022
• Developed real-time analytics platform using Apache Kafka and Elasticsearch
• Designed and implemented REST APIs serving 1M+ requests daily
• Improved system reliability from 99.5% to 99.99% uptime
• Led code review process and established coding standards

Software Engineer - StartUp Labs
January 2018 - June 2019
• Built customer-facing web applications using React and Node.js
• Implemented CI/CD pipelines using Jenkins and Docker
• Collaborated with product team to deliver features on schedule
• Reduced database query time by 60% through optimization

TECHNICAL SKILLS
Languages: Python, Go, JavaScript, TypeScript, Java, SQL
Frameworks: FastAPI, Django, Express.js, React, Spring Boot
Cloud & DevOps: AWS, Google Cloud Platform, Kubernetes, Docker, Terraform
Databases: PostgreSQL, MongoDB, Redis, Elasticsearch
Tools: Git, Jenkins, GitLab CI, Prometheus, Grafana, DataDog

EDUCATION
Master of Science in Computer Science
Tech Institute of Technology
Graduated: May 2017
GPA: 3.9/4.0

Bachelor of Science in Software Engineering
State University
Graduated: May 2015
GPA: 3.7/4.0

CERTIFICATIONS
Certified Kubernetes Application Developer (CKAD)
AWS Certified Solutions Architect - Professional
Google Cloud Professional Data Engineer
"""


def create_docx_resume(output_path):
    """Create a sample resume in DOCX format."""
    doc = Document()
    
    for line in RESUME_CONTENT.split('\n'):
        if line.strip():
            # Add with appropriate formatting
            if line.isupper() and len(line) < 50:
                p = doc.add_paragraph(line, style='Heading 1')
            elif line.startswith('• '):
                p = doc.add_paragraph(line, style='List Bullet')
            else:
                p = doc.add_paragraph(line)
    
    doc.save(output_path)
    print(f"Created {output_path}")


if __name__ == "__main__":
    # Create sample DOCX resume
    docx_path = tests_dir / "sample_resume.docx"
    create_docx_resume(docx_path)
    print(f"Sample DOCX resume created at {docx_path}")
